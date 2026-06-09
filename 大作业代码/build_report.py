# -*- coding: utf-8 -*-
"""Build the complete Word report for the PyTorch breast cancer assignment."""

from __future__ import annotations

import csv
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DOCX = Path(os.getenv("TEMP", "")) / "ml_assignment_template_converted.docx"
OUTPUT_DOCX = PROJECT_ROOT / "0.0大作业完整版报告-乳腺癌分类.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x0B, 0x25, 0x45)
MUTED = RGBColor(0x66, 0x66, 0x66)
BLACK = RGBColor(0x00, 0x00, 0x00)

BASE_FONT = "Calibri"
CN_FONT = "微软雅黑"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def dxa_to_inches(value: int) -> float:
    return value / 1440


def set_run_font(run, name: str = BASE_FONT, cn_name: str = CN_FONT, size: float | None = None,
                 color: RGBColor | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 6,
                          line_spacing: float = 1.10) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    cell.width = Inches(dxa_to_inches(width_dxa))
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_text(cell, text: str, bold: bool = False, color: RGBColor | None = None,
                   align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT, size: float = 9.2) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.05)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color or BLACK, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr_cells[idx], "F2F4F7")
        set_table_text(hdr_cells[idx], header, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_table_header(table.rows[0])

    for row_data in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 and len(value) <= 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_table_text(row_cells[idx], value, align=align)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4)


def paragraph_border_bottom(paragraph, color: str = "2E74B5", size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    field_run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    field_run._r.append(fld_begin)
    field_run._r.append(instr)
    field_run._r.append(fld_end)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BASE_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    def para_style(name: str):
        try:
            return styles[name]
        except KeyError:
            return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = para_style(name)
        style.font.name = BASE_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.keep_with_next = True

    caption = para_style("Caption")
    caption.font.name = BASE_FONT
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    caption.font.size = Pt(9.5)
    caption.font.color.rgb = MUTED
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = ""
    left = header.add_run("机器学习大作业")
    set_run_font(left, size=9, color=MUTED, bold=True)
    right = header.add_run("  |  基于 PyTorch 的乳腺癌良恶性分类")
    set_run_font(right, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    add_page_number(footer)


def add_paragraph(doc: Document, text: str = "", bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        run = p.add_run(text)
        set_run_font(run)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=9.5, color=MUTED)


def add_figure(doc: Document, image_name: str, caption: str, width_in: float = 6.1) -> None:
    image_path = PROJECT_ROOT / "结果图" / image_name
    if not image_path.exists():
        add_paragraph(doc, f"图像文件缺失：{image_name}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    add_caption(doc, caption)


def read_metrics() -> list[dict[str, str]]:
    metrics_path = PROJECT_ROOT / "results" / "model_metrics.csv"
    with metrics_path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def fmt(value: str, digits: int = 4) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except ValueError:
        return value


def build_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    run = p.add_run("机器学习课程大作业报告")
    set_run_font(run, size=16, color=MUTED, bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4, line_spacing=1.05)
    run = p.add_run("基于 PyTorch 的乳腺癌良恶性分类系统设计与实现")
    set_run_font(run, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=16)
    run = p.add_run("任务 2.3：Kaggle 乳腺癌数据集良恶性分类、可视化与模型结果比较")
    set_run_font(run, size=12.5, color=MUTED)

    meta_rows = [
        ("课程名称", "机器学习"),
        ("报告类型", "1.0 完整大作业报告（完整版）"),
        ("小组编号", "__________"),
        ("负责人", "__________"),
        ("成员姓名与学号", "__________"),
        ("成员及工作量分配", "待填写"),
        ("GitHub 仓库", "https://github.com/Esonk/Machine-Learning-Major-Assignment"),
        ("完成日期", "2026 年 6 月 9 日"),
    ]
    for label, value in meta_rows:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=3)
        r1 = p.add_run(f"{label}：")
        set_run_font(r1, size=11, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=11)

    rule = doc.add_paragraph()
    set_paragraph_spacing(rule, before=12, after=10)
    paragraph_border_bottom(rule)

    add_paragraph(
        doc,
        "说明：本报告依据课程完整大作业报告模板撰写，成员信息与工作量分配保留为可编辑占位，提交前可由小组按实际情况补充。"
    )
    doc.add_page_break()


def build_toc(doc: Document) -> None:
    add_heading(doc, "目录", level=1)
    toc_lines = [
        "摘要",
        "1 项目背景与任务说明",
        "2 数据集与预处理",
        "3 模型设计与训练方法",
        "4 工程实现与文件说明",
        "5 实验结果与分析",
        "6 结论、不足与改进方向",
        "参考资料",
        "附录",
    ]
    for line in toc_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        set_paragraph_spacing(p, before=0, after=3, line_spacing=1.10)
        run = p.add_run(line)
        set_run_font(run, size=11)
    doc.add_page_break()


def build_report(doc: Document) -> None:
    metrics = read_metrics()

    build_cover(doc)
    build_toc(doc)

    add_heading(doc, "摘要", level=1)
    add_paragraph(
        doc,
        "本项目面向 Kaggle 平台 Breast Cancer Wisconsin 乳腺癌数据集，完成良性肿瘤与恶性肿瘤的二分类建模、训练、测试和可视化分析。实验使用 PyTorch 完成，围绕 30 个细胞核数值特征构建线性分类器、浅层多层感知机、深层多层感知机和一维卷积网络四类神经网络分类方法，并通过 Accuracy、Precision、Recall、F1-score、Specificity、ROC-AUC 和混淆矩阵等指标比较模型表现。"
    )
    add_paragraph(
        doc,
        "实验结果表明，LinearClassifier、ShallowMLP 与 DeepMLP 在测试集上均达到 0.9882 的 Accuracy 和 0.9841 的 F1-score，其中 DeepMLP 的测试损失最低；FeatureCNN1D 的 Accuracy 为 0.9529，ROC-AUC 为 0.9906。整体来看，该数据集特征维度较低且区分度较强，标准化后的线性模型和 MLP 模型已能取得稳定表现，一维卷积网络可作为特征局部组合模式的探索性对比方法。"
    )
    add_paragraph(doc, "关键词：PyTorch；乳腺癌分类；二分类；多层感知机；一维卷积网络；模型评估", bold_prefix="关键词：")

    add_heading(doc, "1 项目背景与任务说明", level=1)
    add_paragraph(
        doc,
        "乳腺癌良恶性辅助分类是机器学习二分类任务中的典型应用场景。Kaggle 平台提供的 Breast Cancer Wisconsin 数据集包含肿瘤细胞核的半径、纹理、周长、面积、平滑度、凹度、对称性等数值特征，标签表示样本对应肿瘤类型为良性或恶性。本大作业的目标是基于该数据集完成完整的机器学习工程流程：数据读取与清洗、训练集/验证集/测试集划分、特征标准化、四种分类模型训练、评价指标计算、图表可视化以及工程文件整理。"
    )
    add_paragraph(
        doc,
        "项目除 Notebook 实验记录外，还提供一键训练脚本、单样本预测脚本、四个单模型脚本、四个单模型 Notebook、HTML 导出文件、训练权重和结果图。GitHub 仓库用于保存完整源代码、训练数据和实验结果，便于复现和提交检查。"
    )

    add_heading(doc, "2 数据集与预处理", level=1)
    add_heading(doc, "2.1 数据集基本信息", level=2)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["数据集名称", "Breast Cancer Wisconsin"],
            ["数据来源", "Kaggle 数据集：priyanka841/breast-cancer-wisconsin"],
            ["本地数据文件", "数据集/breast cancer.csv"],
            ["原始样本数量", "569 条"],
            ["原始字段数量", "33 列"],
            ["清洗后输入特征", "30 个数值特征"],
            ["分类标签", "diagnosis：B 表示良性，M 表示恶性"],
        ],
        [2600, 6760],
    )
    add_paragraph(
        doc,
        "原始数据中 id 字段仅为样本编号，不参与模型训练；Unnamed: 32 为空列，也不具备建模意义。因此预处理阶段删除这两列，只保留 diagnosis 标签列和 30 个数值特征。标签映射为 B=0、M=1，便于使用神经网络二分类输出。"
    )

    add_heading(doc, "2.2 标签分布与数据划分", level=2)
    add_paragraph(
        doc,
        "清洗后共有 569 条样本，其中良性样本 357 条、恶性样本 212 条。为了保持训练和测试阶段类别比例稳定，实验采用固定随机种子 42 进行分层划分，训练集、验证集、测试集比例为 70%、15%、15%。标准化参数只由训练集计算，验证集和测试集复用训练集均值与标准差，避免把验证或测试信息泄漏到训练阶段。"
    )
    add_table(
        doc,
        ["数据子集", "总样本数", "良性 B", "恶性 M", "用途"],
        [
            ["训练集", "398", "250", "148", "模型参数学习"],
            ["验证集", "86", "54", "32", "早停判断和最优权重选择"],
            ["测试集", "85", "53", "32", "最终泛化性能评估"],
        ],
        [1450, 1350, 1350, 1350, 3860],
    )
    add_figure(doc, "class_distribution.png", "图 1 标签类别分布", width_in=5.3)
    add_paragraph(
        doc,
        "图 1 显示良性样本数量多于恶性样本，但两类样本均具有足够规模。模型评价时不能只看 Accuracy，还需要结合 Recall、Specificity、F1-score 和混淆矩阵观察两类样本的识别情况。"
    )
    add_figure(doc, "feature_correlation_heatmap.png", "图 2 特征相关性热力图", width_in=6.15)
    add_paragraph(
        doc,
        "图 2 表明部分半径、周长、面积相关特征之间具有较强相关性，说明数据中存在冗余信息。神经网络模型可以在标准化输入基础上学习特征权重与组合关系，但过深模型也可能在小样本表格数据上引入额外波动，因此本实验同时比较简单模型和复杂模型。"
    )

    add_heading(doc, "3 模型设计与训练方法", level=1)
    add_heading(doc, "3.1 四种分类模型", level=2)
    add_table(
        doc,
        ["模型", "结构设计", "设计目的"],
        [
            ["LinearClassifier", "30 维输入直接连接到 2 维输出", "作为最简单基线，检验线性可分程度"],
            ["ShallowMLP", "Linear(30,32) + ReLU + Linear(32,2)", "通过一层隐藏层学习非线性特征组合"],
            ["DeepMLP", "30→64→32→16→2，含 BatchNorm 和 Dropout", "提升表达能力并通过正则化控制过拟合"],
            ["FeatureCNN1D", "将 30 个特征视为一维序列，Conv1d(1,16) 和 Conv1d(16,32) 后分类", "探索相邻特征局部组合对分类的影响"],
        ],
        [1700, 4000, 3660],
    )
    add_paragraph(
        doc,
        "四个模型的输出维度均为 2，对应良性和恶性两个类别。训练阶段使用 CrossEntropyLoss，该损失函数内部完成类别 logits 到概率分布的处理；预测阶段再通过 softmax 得到良性概率和恶性概率。"
    )

    add_heading(doc, "3.2 训练配置与评价指标", level=2)
    add_table(
        doc,
        ["配置项", "设置"],
        [
            ["随机种子", "42"],
            ["优化器", "Adam"],
            ["学习率", "1e-3"],
            ["权重衰减", "1e-4"],
            ["批量大小", "32"],
            ["最大训练轮数", "300"],
            ["早停策略", "验证集 loss 连续 30 轮无提升则停止"],
            ["权重保存", "保存验证集 F1-score 最优的模型权重"],
        ],
        [2800, 6560],
    )
    add_paragraph(
        doc,
        "评价指标包括 Accuracy、Precision、Recall、Specificity、F1-score、ROC-AUC 和 Confusion Matrix。其中 Recall 关注恶性样本被正确识别的比例，Specificity 关注良性样本被正确识别的比例，F1-score 综合 Precision 与 Recall，ROC-AUC 衡量不同阈值下模型区分类别的整体能力。"
    )

    add_heading(doc, "4 工程实现与文件说明", level=1)
    add_paragraph(
        doc,
        "工程目录独立放置在 D:\\桌面\\机器学习大作业\\一.大作业\\2代码\\2.3乳腺癌分类。目录结构按课程提交要求整理，包含数据、Notebook、Python 源码、HTML 导出、结果图、训练权重、指标表和说明文档。"
    )
    add_table(
        doc,
        ["文件或目录", "说明"],
        [
            ["数据集/breast cancer.csv", "Kaggle 下载后的完整训练数据"],
            ["大作业代码/breast_cancer_pytorch.ipynb", "完整实验流程 Notebook，包含 Markdown 分析、代码和输出"],
            ["大作业代码/LinearClassifier.py 等四个脚本", "四个模型分别对应的独立训练脚本"],
            ["大作业代码/LinearClassifier（线性分类器）.ipynb 等四个 Notebook", "按模型拆分的 Notebook 实验文件"],
            ["大作业代码/train_models.py", "一键训练四个模型并保存结果"],
            ["大作业代码/predict_one.py", "加载最佳模型并完成单样本预测"],
            ["html格式/*.html", "Notebook 导出的 HTML 结果文件"],
            ["结果图/*.png", "类别分布、热力图、训练曲线、混淆矩阵、ROC 曲线和模型对比图"],
            ["results/*.pt / *.csv / *.json", "模型权重、训练历史、指标表和标准化参数"],
            ["README.md / requirements.txt", "项目说明和运行依赖"],
        ],
        [3600, 5760],
    )
    add_paragraph(
        doc,
        "GitHub 仓库链接：https://github.com/Esonk/Machine-Learning-Major-Assignment。提交报告时可将该链接填写到完整版报告的对应位置，用于展示完整源代码和数据文件。"
    )

    add_heading(doc, "5 实验结果与分析", level=1)
    add_heading(doc, "5.1 训练过程", level=2)
    add_figure(doc, "training_curves.png", "图 3 四个模型训练与验证曲线", width_in=6.25)
    add_paragraph(
        doc,
        "图 3 展示了四个模型训练过程中 loss 和主要指标的变化。数据规模较小，四个模型均能在较少训练轮数内收敛；早停策略避免模型在验证集性能不再提升后继续训练。LinearClassifier 的收敛轮次较多，说明简单线性模型需要更长训练过程达到稳定；ShallowMLP 和 DeepMLP 较快达到较优验证结果。"
    )

    add_heading(doc, "5.2 测试集指标对比", level=2)
    metric_rows = []
    for row in metrics:
        metric_rows.append([
            row["model"],
            row["best_epoch"],
            fmt(row["test_loss"]),
            fmt(row["accuracy"]),
            fmt(row["precision"]),
            fmt(row["recall"]),
            fmt(row["specificity"]),
            fmt(row["f1"]),
            fmt(row["roc_auc"]),
        ])
    add_table(
        doc,
        ["模型", "最佳轮次", "测试损失", "Accuracy", "Precision", "Recall", "Specificity", "F1", "ROC-AUC"],
        metric_rows,
        [1700, 900, 1000, 950, 950, 900, 1100, 850, 1010],
    )
    add_paragraph(
        doc,
        "从指标表可见，LinearClassifier、ShallowMLP 和 DeepMLP 在测试集上均只漏判 1 个恶性样本，没有将良性样本误判为恶性，因此 Accuracy、Precision、Recall、Specificity、F1-score 和 ROC-AUC 完全一致。DeepMLP 的测试损失最低，说明它在正确分类基础上输出概率更稳定。FeatureCNN1D 在测试集中出现 2 个假阳性和 2 个假阴性，Accuracy 和 F1-score 低于其他三个模型，但 ROC-AUC 仍达到 0.9906，说明其排序区分能力仍较强。"
    )
    add_figure(doc, "confusion_matrices.png", "图 4 四个模型测试集混淆矩阵", width_in=6.05)
    add_paragraph(
        doc,
        "图 4 进一步说明，前三个模型的错误集中在 1 个恶性样本漏判，FeatureCNN1D 同时存在良性误判为恶性和恶性误判为良性的情况。对于良恶性二分类任务，混淆矩阵能够比单一准确率更直观地展示错误类型。"
    )
    add_figure(doc, "roc_curves.png", "图 5 四个模型 ROC 曲线", width_in=6.05)
    add_paragraph(
        doc,
        "图 5 显示四个模型的 ROC 曲线均靠近左上角，说明它们在不同阈值下都具有较好的区分能力。其中前三个模型测试 ROC-AUC 为 1.0000，FeatureCNN1D 为 0.9906。考虑到测试集规模为 85 条，AUC 差异需要结合混淆矩阵和多次实验进一步判断，不能简单外推为真实临床场景中的绝对性能。"
    )
    add_figure(doc, "model_comparison.png", "图 6 四个模型主要指标对比", width_in=6.05)
    add_paragraph(
        doc,
        "图 6 对比了各模型的主要测试指标。整体结论是：在该数据集上，线性模型已经具备很强的分类能力；浅层和深层 MLP 在保持高性能的同时提供了非线性建模能力；一维卷积网络可用于探索特征局部组合，但对表格数据的优势并不明显。"
    )

    add_heading(doc, "6 结论、不足与改进方向", level=1)
    add_paragraph(
        doc,
        "本项目完成了基于 PyTorch 的乳腺癌良恶性分类实验。完整流程包括数据清洗、分层划分、标准化、四种神经网络模型训练、手写评价指标计算、可视化分析、模型权重保存和单样本预测程序。实验结果显示，LinearClassifier、ShallowMLP 和 DeepMLP 在当前测试集上取得相同的分类指标，DeepMLP 的测试损失最低；FeatureCNN1D 表现略低，但仍具有较高 ROC-AUC。"
    )
    add_paragraph(
        doc,
        "本实验的不足主要包括：数据集规模较小，单次训练/验证/测试划分可能带来偶然性；表格数据中特征顺序并不一定具有真实空间邻近关系，因此一维卷积网络的结构假设需要谨慎解释；本项目重点是课程实验和工程复现，不能直接用于医学诊断决策。后续可进一步加入多随机种子重复实验、交叉验证思想、特征重要性分析、阈值敏感性分析和更规范的模型校准评估。"
    )

    add_heading(doc, "参考资料", level=1)
    add_paragraph(doc, "[1] Kaggle. Breast Cancer Wisconsin Dataset. https://www.kaggle.com/datasets/priyanka841/breast-cancer-wisconsin")
    add_paragraph(doc, "[2] PyTorch Documentation. https://pytorch.org/docs/")
    add_paragraph(doc, "[3] UCI Machine Learning Repository. Breast Cancer Wisconsin Diagnostic Dataset.")

    add_heading(doc, "附录", level=1)
    add_heading(doc, "附录 A 运行环境", level=2)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["操作系统", "Windows"],
            ["Python 环境", "Anaconda 环境 geo3d"],
            ["主要依赖", "torch、pandas、numpy、matplotlib、seaborn、notebook、nbconvert"],
            ["训练设备", "CPU"],
            ["运行入口", "大作业代码/train_models.py"],
        ],
        [2600, 6760],
    )
    add_heading(doc, "附录 B 提交材料检查", level=2)
    add_table(
        doc,
        ["检查项", "状态"],
        [
            ["完整训练数据", "已包含"],
            ["完整 Notebook 与输出", "已包含"],
            ["HTML 格式 Notebook", "已导出"],
            ["四个模型方法", "已实现并比较"],
            ["结果图与指标表", "已保存"],
            ["GitHub 仓库链接", "已写入报告"],
            ["成员与工作量信息", "保留占位，提交前可补充"],
        ],
        [3600, 5760],
    )


def main() -> None:
    if TEMPLATE_DOCX.exists():
        doc = Document(str(TEMPLATE_DOCX))
        clear_document_body(doc)
    else:
        doc = Document()
    configure_document(doc)
    build_report(doc)

    doc.core_properties.title = "基于 PyTorch 的乳腺癌良恶性分类系统设计与实现"
    doc.core_properties.subject = "机器学习大作业完整报告"
    doc.core_properties.author = "__________"
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
